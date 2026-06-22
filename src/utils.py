import io
from docx import Document
import pdfplumber
from fastapi import UploadFile
import os 


async def read_document(file: UploadFile) -> str:
    """Read text from DOCX, PDF, or TXT files"""
    raw = await file.read()
    ext = os.path.splitext(file.filename)[1].lower()

    if ext == '.docx':
        return read_docx(raw)
    elif ext == '.pdf':
        return read_pdf(raw)
    elif ext == '.txt':
        return read_txt(raw)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def read_txt(raw: bytes) -> str:
    """Read text content from a TXT file"""
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1")


def read_docx(raw: bytes) -> str:
    """Read text content from a DOCX file"""
    doc = Document(io.BytesIO(raw))
    full_text = []

    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)


def read_pdf(raw: bytes) -> str:
    """Read text from PDF using pdfplumber"""
    text = ""
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

import re

def parse_evaluation(text):
    """
    Parses a string containing a score and a 'why' explanation 
    into a dictionary with keys 'score' and 'reason'.
    """
    result = {'score': None, 'reason': None}
    
    # 1. Extract the score
    # Looks for "score", optional spaces, a colon, optional quotes, and captures the digits
    score_match = re.search(r'score\s*:\s*[\'"]?(\d+)[\'"]?', text, re.IGNORECASE)
    if score_match:
        # Casts the extracted score to an integer
        result['score'] = int(score_match.group(1))
        
    # 2. Extract the reason (why)
    # Looks for "why", optional spaces, a colon, optional starting quotes, 
    # and captures everything else until the ending quote/end of string.
    # re.DOTALL ensures it captures the text even if it spans multiple lines.
    why_match = re.search(r'why\s*:\s*[\'"]?(.*?)[\'"]?\s*$', text, re.IGNORECASE | re.DOTALL)
    if why_match:
        # Extracts the matched string
        result['reason'] = why_match.group(1)
        
    return result
from docx import Document
import pdfplumber
from fastapi import UploadFile
import os 


async def read_document(file: UploadFile) -> str:
    """Read text from DOCX, PDF, or TXT files"""
    raw = await file.read()
    ext = os.path.splitext(file.filename)[1].lower()

    if ext == '.docx':
        return read_docx(raw)
    elif ext == '.pdf':
        return read_pdf(raw)
    elif ext == '.txt':
        return read_txt(raw)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def read_txt(raw: bytes) -> str:
    """Read text content from a TXT file"""
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1")


def read_docx(raw: bytes) -> str:
    """Read text content from a DOCX file"""
    doc = Document(io.BytesIO(raw))
    full_text = []

    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)


def read_pdf(raw: bytes) -> str:
    """Read text from PDF using pdfplumber"""
    text = ""
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

import re

def parse_evaluation(text):
    """
    Parses a string containing a score and a 'why' explanation 
    into a dictionary with keys 'score' and 'reason'.
    """
    result = {'score': None, 'reason': None}
    
    # 1. Extract the score
    # Looks for "score", optional spaces, a colon, optional quotes, and captures the digits
    score_match = re.search(r'score\s*:\s*[\'"]?(\d+)[\'"]?', text, re.IGNORECASE)
    if score_match:
        # Casts the extracted score to an integer
        result['score'] = int(score_match.group(1))
        
    # 2. Extract the reason (why)
    # Looks for "why", optional spaces, a colon, optional starting quotes, 
    # and captures everything else until the ending quote/end of string.
    # re.DOTALL ensures it captures the text even if it spans multiple lines.
    why_match = re.search(r'why\s*:\s*[\'"]?(.*?)[\'"]?\s*$', text, re.IGNORECASE | re.DOTALL)
    if why_match:
        # Extracts the matched string
        result['reason'] = why_match.group(1)
        
    return result
from docx import Document
import pdfplumber
from fastapi import UploadFile
import os 


async def read_document(file: UploadFile) -> str:
    """Read text from DOCX, PDF, or TXT files"""
    raw = await file.read()
    ext = os.path.splitext(file.filename)[1].lower()

    if ext == '.docx':
        return read_docx(raw)
    elif ext == '.pdf':
        return read_pdf(raw)
    elif ext == '.txt':
        return read_txt(raw)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def read_txt(raw: bytes) -> str:
    """Read text content from a TXT file"""
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1")


def read_docx(raw: bytes) -> str:
    """Read text content from a DOCX file"""
    doc = Document(io.BytesIO(raw))
    full_text = []

    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)


def read_pdf(raw: bytes) -> str:
    """Read text from PDF using pdfplumber"""
    text = ""
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

